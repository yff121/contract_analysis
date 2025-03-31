import os
import tempfile
import streamlit as st
import requests
from docx import Document
from io import BytesIO
import pdfplumber

# 第一部分：法规文件处理
class RegulationLoader:
    def __init__(self):
        self.regulations = []

    def load_documents(self, uploaded_files):
        """处理上传的文件对象"""
        for file in uploaded_files:
            filename = file.name
            
            if filename.endswith('.pdf'):
                content = self._read_pdf(file)
            elif filename.endswith('.docx'):
                content = self._read_docx(file)
            else:
                continue
            
            if content:
                self.regulations.append({
                    'filename': filename,
                    'content': content
                })

    def _read_pdf(self, file):
        """读取PDF文件内容"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.getvalue())
                tmp_path = tmp.name
            
            content = []
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    content.append(page.extract_text())
            os.unlink(tmp_path)
            return '\n'.join(content)
        except Exception as e:
            st.error(f"PDF读取失败 {file.name}: {str(e)}")
            return None

    def _read_docx(self, file):
        """读取Word文档内容"""
        try:
            doc = Document(BytesIO(file.getvalue()))
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            st.error(f"DOCX读取失败 {file.name}: {str(e)}")
            return None

# 第二部分：合规性分析
class ComplianceChecker:
    def __init__(self, api_key):
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.api_key = api_key
        self.headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }

    def analyze_compliance(self, contract_text, regulation):
        """使用DeepSeek API进行合规性分析"""
        prompt = self._build_prompt(contract_text, regulation)
        
        payload = {
            "model": "deepseek-reasoner",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.2
        }

        try:
            response = requests.post(
                self.api_url,
                headers=self.headers,
                json=payload
            )
            response.raise_for_status()
            return response.json()['choices'][0]['message']['content']
        except Exception as e:
            st.error(f"API请求失败: {str(e)}")
            return None

    def _build_prompt(self, contract_text, regulation):
        """构建分析提示词"""
        return f"""请进行合规性分析，严格按照以下要求执行：
        1. 分析以下合同内容是否符合 {regulation['filename']} 的要求，合同中的内容与 {regulation['filename']}中的不冲突就算符合要求；
        2. 合同相对 {regulation['filename']}会有缺失的内容，但是这不算有冲突；
        3. 如果确实存在内容有冲突，请指出具体不符合的条款，并给出修改建议；
        4. 依据规定，指出合同还应该增加的内容。

        合同内容：
        {contract_text[:30000]}

        相关法规内容：
        {regulation['content'][:30000]}

        请用规范的报告格式返回分析结果，包含：分析的法规名称、合规状态、问题点与修改建议（内容以表格的形式呈现）、还应增加的内容（内容以表格的形式呈现）。
        不用输出报告单位和日期。
        """

# Streamlit界面
st.set_page_config(layout="wide")
st.title("📑 智能合同合规性检查系统")
st.markdown("""
<style>
    .stApp {
        background-color: #f5f7fa;
    }
    .stButton>button {
        background-color: #4a6bdf;
        color: white;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 500;
    }
    .stFileUploader>div>div>div>div {
        border: 2px dashed #4a6bdf;
        border-radius: 8px;
        padding: 2rem;
    }
    .stExpander {
        background-color: white;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
</style>
""", unsafe_allow_html=True)

# 初始化session state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = []

# 侧边栏配置
with st.sidebar:
    st.header("⚙️ 配置参数")
    # 从session state获取API Key，如果没有则显示空
    api_key = st.text_input(
        "🔑 DeepSeek API Key", 
        value=st.session_state.get("api_key", ""),
        type="password", 
        help="请从DeepSeek官网获取API密钥"
    )
    # 将输入的API Key保存到session state
    st.session_state["api_key"] = api_key
    st.info("请使用您自己的API密钥", icon="ℹ️")

# 文件上传区域
st.subheader("📤 上传文件")
col1, col2 = st.columns(2, gap="large")
with col1:
    uploaded_regs = st.file_uploader(
        "上传法规文件（PDF/DOCX）",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        help="可上传多个法规文件"
    )
with col2:
    uploaded_contract = st.file_uploader(
        "上传合同文件（DOCX）",
        type=["docx"],
        help="请上传需要检查的合同文件"
    )

# 显示历史分析结果
if st.session_state.analysis_results:
    st.divider()
    st.subheader("📊 分析结果")
    for idx, result in enumerate(st.session_state.analysis_results, 1):
        with st.expander(f"📜 {result['filename']} 分析结果", expanded=False):
            st.markdown(result['content'])

# 分析处理逻辑
if st.button("🚀 开始分析", use_container_width=True, type="primary"):
    if not uploaded_regs:
        st.warning("⚠️ 请先上传法规文件", icon="⚠️")
        st.stop()
    if not uploaded_contract:
        st.warning("⚠️ 请先上传合同文件", icon="⚠️")
        st.stop()
    if not api_key:
        st.warning("⚠️ 请输入API密钥", icon="⚠️")
        st.stop()

    with st.spinner("⏳ 正在初始化..."):
        # 加载合同内容
        try:
            contract_doc = Document(BytesIO(uploaded_contract.getvalue()))
            contract_text = '\n'.join([para.text for para in contract_doc.paragraphs if para.text])
        except Exception as e:
            st.error(f"❌ 合同读取失败: {str(e)}")
            st.stop()

        # 加载法规文件
        loader = RegulationLoader()
        loader.load_documents(uploaded_regs)
        
        if not loader.regulations:
            st.error("❌ 没有成功加载任何法规文件")
            st.stop()

        # 初始化检查器
        checker = ComplianceChecker(api_key)

    # 清空历史结果
    st.session_state.analysis_results = []

    # 分析处理
    progress_bar = st.progress(0)
    total_regs = len(loader.regulations)

    for idx, regulation in enumerate(loader.regulations):
        progress = (idx + 1) / total_regs
        progress_bar.progress(progress, text=f"📝 正在分析法规 {idx+1}/{total_regs}")
        
        with st.expander(f"📜 {regulation['filename']} 分析结果", expanded=False):
            try:
                result = checker.analyze_compliance(contract_text, regulation)
                if result:
                    st.markdown(result)
                    st.session_state.analysis_results.append({
                        'filename': regulation['filename'],
                        'content': result
                    })
                else:
                    st.error("❌ 分析失败")
            except Exception as e:
                st.error(f"❌ 分析过程中发生错误: {str(e)}")
    
    st.balloons()
    st.success("✅ 分析完成！")
    progress_bar.empty()